"""
Script para preencher o formulário de registro de software (UESC/NIT)
com as informações do projeto Química RA.
"""
from docx import Document
from docx.shared import Pt
import copy, re

SRC  = r"e:\PROJECT\MILI\assets\formulario-registro-software.docx"
DEST = r"e:\PROJECT\MILI\assets\formulario-registro-software-PREENCHIDO.docx"

# ─────────────────────────────────────────
#  Conteúdo do projeto
# ─────────────────────────────────────────
TITULO = (
    "Química RA: Aplicativo Educacional de Realidade Aumentada "
    "para o Ensino de Química Atmosférica"
)
DATA_CONCLUSAO  = "2024"
LINGUAGENS      = "TypeScript e JavaScript (React Native / Expo SDK)"
DEPENDENCIAS    = (
    "Sim. Principais dependências: Expo SDK 53.0.19, React Native 0.79.5, "
    "@reactvision/react-viro 2.43.3, expo-camera 16.1.10, expo-gl 15.1.7, "
    "React 19.0.0, TypeScript 5.x."
)
PUBLICO_ALVO = (
    "Estudantes do Ensino Médio e Superior nas áreas de Química e Ciências, "
    "professores de química e educadores científicos."
)
RESUMO = (
    "O Química RA é um aplicativo educacional para dispositivos Android que utiliza "
    "Realidade Aumentada (RA) para ensinar conceitos de química atmosférica de forma "
    "interativa e imersiva. O aplicativo permite que estudantes visualizem modelos 3D "
    "de compostos químicos e moléculas presentes na atmosfera terrestre diretamente "
    "através da câmera do celular, em ambiente de RA. Está organizado em quatro "
    "módulos temáticos: (1) Composição Atmosférica – N₂, O₂ e Ar; (2) Compostos "
    "Químicos e Seus Impactos – SO₂ e NO₂; (3) Efeitos na Atmosfera – dividido em "
    "Efeito Estufa (CO₂, CH₄, N₂O, H₂O vapor) e Camada de Ozônio (O₃, CFC-11); "
    "(4) Preservação e Possíveis Soluções – módulo textual com estratégias de mitigação. "
    "O aplicativo integra visualização 3D interativa com controles adaptativos de "
    "rotação e zoom, compatíveis com diferentes fabricantes de smartphones Android "
    "(incluindo tratamento especial para dispositivos Xiaomi/MIUI)."
)
DESCRICAO_DETALHADA = (
    "O aplicativo Química RA foi desenvolvido utilizando o framework React Native "
    "com Expo SDK 53.0.19 e a linguagem TypeScript, destinado à plataforma Android. "
    "O núcleo de Realidade Aumentada é implementado com o SDK @reactvision/react-viro "
    "2.43.3, que fornece componentes como ViroARSceneNavigator, ViroARScene, ViroNode, "
    "Viro3DObject e ViroAmbientLight para renderização de modelos moleculares 3D "
    "em ambiente AR. A navegação é gerenciada pelo Expo Router com roteamento baseado "
    "em sistema de arquivos.\n"
    "O sistema é estruturado em quatro módulos educacionais (capítulos):\n"
    "1. Composição Atmosférica: Apresenta os principais gases da atmosfera – N₂ "
    "(nitrogênio, 78%, geometria linear, apolar, ligação tripla), O₂ (oxigênio, 21%, "
    "geometria linear, apolar, ligação dupla) e Ar (argônio, 0,093%, gás nobre inerte) "
    "– com visualização 3D e informações científicas associadas.\n"
    "2. Compostos Químicos e Seus Impactos: Aborda compostos poluentes atmosféricos: "
    "SO₂ (dióxido de enxofre, geometria angular, 119°, polar) e NO₂ (dióxido de "
    "nitrogênio, geometria angular, 134°, polar), relacionando-os com suas origens e "
    "impactos ambientais.\n"
    "3. Efeitos na Atmosfera – subdividido em: (3a) Efeito Estufa: CO₂ (linear, apolar, "
    "180°), CH₄ (tetraédrica, apolar, 109,5°), N₂O (linear, levemente polar) e H₂O "
    "vapor (angular, polar, 104,5°); (3b) Camada de Ozônio: O₃ (angular, polar, 116,8°) "
    "e CFC-11 CCl₃F (tetraédrica, polar, 109,5°).\n"
    "4. Preservação e Possíveis Soluções: Módulo textual com estratégias de mitigação "
    "de emissões poluentes, abordando fontes de energia limpas, transportes sustentáveis, "
    "reflorestamento, créditos de carbono e políticas públicas.\n"
    "O sistema de controles é adaptativo: em dispositivos Xiaomi/MIUI, são exibidos "
    "botões direcionais para rotação manual (±15° por eixo X e Y); em outros "
    "dispositivos Android, são utilizados gestos nativos via PanResponder (pinça para "
    "zoom, arrastar com 2 dedos para rotação livre). Os modelos 3D são armazenados no "
    "formato GLB e carregados via sistema de registro de modelos (modelRegistry). "
    "A interface utiliza tema escuro com paleta de cores consistente, componentes "
    "LinearGradient, ScrollView e ImageBackground."
)
DATA_PUBLICACAO = "2024"
APLICACOES = (
    "O Química RA destina-se ao ensino de química atmosférica no Ensino Médio e "
    "Superior. As principais aplicações incluem:\n"
    "1. Material Didático Complementar: Complementa o ensino teórico de química "
    "tornando conceitos abstratos – como geometria molecular, polaridade e ângulos de "
    "ligação – mais tangíveis por meio de visualização 3D imersiva em Realidade Aumentada.\n"
    "2. Visualização de Compostos Atmosféricos: Permite ao estudante visualizar e "
    "interagir com modelos moleculares 3D de compostos presentes na atmosfera "
    "(N₂, O₂, Ar, SO₂, NO₂, CO₂, CH₄, N₂O, H₂O, O₃, CFC-11).\n"
    "3. Ensino sobre Impactos Ambientais: Contextualiza compostos químicos com seus "
    "efeitos ambientais: poluição do ar, efeito estufa e destruição da camada de ozônio.\n"
    "4. Apoio a Pesquisas Educacionais: Ferramenta para pesquisas sobre o uso de "
    "Realidade Aumentada (RA) no ensino de ciências, avaliando impacto no aprendizado.\n"
    "5. Formação de Professores: Pode compor material de capacitação sobre uso de "
    "tecnologias educacionais inovadoras (RA) nas aulas de química e ciências."
)
# Tipos selecionados (marcados com ✓ no início do parágrafo)
# Tipos (seção "Tipo do Programa de Computador", paras 147-249)
TIPOS_SELECIONADOS = {
    "AP01",   # Aplicativos
    "SM04",   # CAE/CAD/CAM (inclui CBT - Computer-Based Training)
    "TC01",   # Aplicações Técnico-Científicas
}
# Áreas de aplicação (seção "Aplicação do Programa de Computador", paras 256+)
AREAS_SELECIONADAS = {
    "ED01",   # Ensino regular (2º grau e superior)
    "ED04",   # Formas de ensino / material instrucional (audiovisual, didático)
    "FQ14",   # Química (composto químico, substância química)
}
PATROCINADORES = (
    "Projeto desenvolvido no âmbito de Trabalho de Conclusão de Curso (Mestrado) "
    "na Universidade Estadual de Santa Cruz (UESC), Ilhéus – BA. "
    "Não há envolvimento de patrocinadores externos ou convênios com empresas."
)


# ─────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────
def clear_para(para):
    """Remove todos os runs de um parágrafo."""
    for run in list(para.runs):
        run.text = ""

def set_para_text(para, text, bold=False):
    """Limpa o parágrafo e insere um único run com o texto."""
    clear_para(para)
    run = para.add_run(text)
    run.bold = bold

def replace_run_content(para, label_run_count, new_text):
    """
    Mantém os primeiros `label_run_count` runs (rótulo em negrito)
    e substitui o restante por um único run com new_text.
    """
    # Apagar runs após o rótulo
    runs = para.runs
    for i in range(label_run_count, len(runs)):
        runs[i].text = ""
    # Adicionar novo run com o conteúdo
    run = para.add_run(new_text)
    run.bold = False

def mark_tipo(para, codigos_selecionados):
    """
    Se o parágrafo começa com um código de tipo (ex.: 'AP01' ou 'AP01-'),
    adiciona '✓ ' na frente se o código estiver selecionado.
    """
    text = para.text.strip()
    # Extrai código: 2 letras + 2 dígitos no início da linha
    match = re.match(r'^([A-Z]{2}\d{2})', text)
    if match:
        code = match.group(1)
        if code in codigos_selecionados:
            if para.runs:
                para.runs[0].text = "✓ " + para.runs[0].text
            else:
                para.add_run("✓ " + text)


# ─────────────────────────────────────────
#  Principal
# ─────────────────────────────────────────
doc = Document(SRC)
paras = doc.paragraphs

# 1. Data do documento (parágrafo 1)
set_para_text(paras[1], "Ilhéus-BA, 2024.")

# 2. Título (62 = rótulo, 63 = conteúdo)
set_para_text(paras[63], TITULO)

# 3. Data de Conclusão (64 = rótulo, 65 = conteúdo)
set_para_text(paras[65], DATA_CONCLUSAO)

# 4. Linguagens (66 = rótulo, 67 = conteúdo)
set_para_text(paras[67], LINGUAGENS)

# 5. Dependências (68 = label+resposta embutida)
#    run[8] = "dependências? " , run[9] = "Não" → trocar run[9]
paras[68].runs[9].text = "Sim"
# Adicionar detalhe no parágrafo 69 (vazio após)
set_para_text(paras[69], DEPENDENCIAS)

# 6. Derivação (70) – já "Não", manter

# 7. Público-alvo (72 = rótulo, 73-74 = vazio)
set_para_text(paras[73], PUBLICO_ALVO)

# 8. Resumo (75 = label+conteúdo antigo)
#    run[0] = bold "Resumo", runs[1..] = conteúdo antigo
replace_run_content(paras[75], 1, ": " + RESUMO)
# Limpar parágrafos 76-84 (itens do resumo antigo)
for i in range(76, 85):
    clear_para(paras[i])

# 9. Descrição Detalhada (87 = label+conteúdo antigo)
#    run[0] = bold "Descrição Detalhada", runs[1..] = conteúdo antigo
replace_run_content(paras[87], 1, ": " + DESCRICAO_DETALHADA)
# Limpar parágrafos 88-98 (conteúdo antigo da descrição)
for i in range(88, 99):
    clear_para(paras[i])

# 10. Data de publicação (100 = label, conteúdo inline)
paras[100].runs[2].text = "de publicação: " + DATA_PUBLICACAO

# 11. Aplicações (101...) – substituir título e bloco antigo
set_para_text(paras[101], "1.10. Aplicações do Programa de Computador:")
set_para_text(paras[104], APLICACOES)
# Limpar parágrafos 105-140 (conteúdo antigo de aplicações)
for i in range(105, 145):
    if i < len(paras):
        clear_para(paras[i])

# 12a. Marcar tipos de programa (seção paras 147–249)
for i in range(147, min(250, len(paras))):
    mark_tipo(paras[i], TIPOS_SELECIONADOS)

# 12b. Marcar áreas de aplicação (seção paras 256 em diante)
for i in range(256, len(paras)):
    mark_tipo(paras[i], AREAS_SELECIONADAS)

# 13. Patrocinadores – preencher campo "( ) Outros. Se sim, especifique:"
for i, para in enumerate(paras):
    if "Outros. Se sim, especifique" in para.text:
        set_para_text(para,
            "(✓) Outros. Se sim, especifique: " + PATROCINADORES)
        break

# ─────────────────────────────────────────
#  Salvar
# ─────────────────────────────────────────
doc.save(DEST)
print(f"Formulário preenchido salvo em:\n  {DEST}")
